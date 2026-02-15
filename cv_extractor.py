# import re
# import pdfplumber
# import docx
# from schemas import CVExtractedData, EducationItem, ExperienceItem


# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     """Extract text from PDF"""
#     import io
#     text = ""
#     with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text.strip()


# def extract_text_from_docx(file_bytes: bytes) -> str:
#     import io
#     doc = docx.Document(io.BytesIO(file_bytes))
#     text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#     return text.strip()


# def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
#     filename_lower = filename.lower()
#     if filename_lower.endswith(".pdf"):
#         return extract_text_from_pdf(file_bytes)
#     elif filename_lower.endswith(".docx"):
#         return extract_text_from_docx(file_bytes)
#     else:
#         raise ValueError("Only PDF or DOCX is supported")


# def extract_email(text: str) -> str:
#     """Extract email"""
#     pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
#     match = re.search(pattern, text)
#     return match.group(0) if match else None


# def extract_phone(text: str) -> str:
#     """Extract phone number"""
#     patterns = [
#         r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
#         r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
#         r'\d{10,}',
#     ]
#     for pattern in patterns:
#         match = re.search(pattern, text)
#         if match:
#             return match.group(0)
#     return None


# def extract_name(text: str) -> str:
    
#     lines = text.split('\n')[:5]  # First 5 lines
#     for line in lines:
#         line = line.strip()
#         # If the line is not too short and doesn't contain email/phone
#         if len(line) > 2 and len(line) < 50:
#             if '@' not in line and not re.search(r'\d{3,}', line):
#                 # Check if it contains only letters and spaces
#                 if re.match(r'^[A-Za-z\s.]+$', line):
#                     return line
#     return None


# def extract_skills(text: str) -> list:
#     """
#     Find common skills.
#     You can add more skills to this list.
#     """
#     skills_keywords = [
#         # Programming
#         'Python', 'Java', 'JavaScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust',
#         'TypeScript', 'Swift', 'Kotlin', 'R', 'MATLAB', 'Scala',
        
#         # Web
#         'HTML', 'CSS', 'React', 'Angular', 'Vue', 'Node.js', 'Express',
#         'Django', 'Flask', 'FastAPI', 'Spring', 'ASP.NET', 'Laravel',
        
#         # Database
#         'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle',
#         'SQLite', 'Cassandra', 'DynamoDB',
        
#         # Tools
#         'Git', 'Docker', 'Kubernetes', 'Jenkins', 'AWS', 'Azure', 'GCP',
#         'Linux', 'REST API', 'GraphQL', 'Postman',
        
#         # Other
#         'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch',
#         'Data Analysis', 'Excel', 'Power BI', 'Tableau',
#     ]
    
#     found_skills = []
#     text_upper = text.upper()
    
#     for skill in skills_keywords:
#         # Case-insensitive search
#         if skill.upper() in text_upper:
#             if skill not in found_skills:
#                 found_skills.append(skill)
    
#     return found_skills


# def extract_education(text: str) -> list:
#     """Extract education"""
#     education = []
    
#     # Find degrees
#     degree_patterns = [
#         r'(B\.?Sc\.?|Bachelor|BSc|BS)\s+(in\s+)?([A-Za-z\s]+)',
#         r'(M\.?Sc\.?|Master|MSc|MS)\s+(in\s+)?([A-Za-z\s]+)',
#         r'(PhD|Ph\.D\.?|Doctorate)\s+(in\s+)?([A-Za-z\s]+)',
#     ]
    
#     for pattern in degree_patterns:
#         matches = re.finditer(pattern, text, re.IGNORECASE)
#         for match in matches:
#             degree = match.group(0).strip()
            
#             # Find year (near the degree)
#             year_match = re.search(r'(19|20)\d{2}', text[max(0, match.start()-50):match.end()+50])
#             year = year_match.group(0) if year_match else None
            
#             education.append(EducationItem(
#                 degree=degree,
#                 institution=None,  # Institution detection is difficult
#                 year=year
#             ))
    
#     return education if education else None


# def extract_experience(text: str) -> list:
#     """Extract work experience"""
#     experience = []
    
#     # Common job titles
#     job_titles = [
#         'Developer', 'Engineer', 'Manager', 'Designer', 'Analyst',
#         'Consultant', 'Specialist', 'Architect', 'Lead', 'Senior',
#         'Junior', 'Intern', 'Associate', 'Director', 'CEO', 'CTO'
#     ]
    
#     for title in job_titles:
#         pattern = rf'\b({title})\b'
#         matches = re.finditer(pattern, text, re.IGNORECASE)
        
#         for match in matches:
#             # Find duration
#             duration_match = re.search(r'(\d+)\s*(year|yr|month|mo)', 
#                                       text[max(0, match.start()-100):match.end()+100], 
#                                       re.IGNORECASE)
#             duration = duration_match.group(0) if duration_match else None
            
#             experience.append(ExperienceItem(
#                 role=match.group(0),
#                 company=None,  # Company name detection is difficult
#                 duration=duration
#             ))
    
#     # Remove duplicates
#     seen = set()
#     unique_exp = []
#     for exp in experience:
#         if exp.role not in seen:
#             seen.add(exp.role)
#             unique_exp.append(exp)
    
#     return unique_exp if unique_exp else None


# def extract_cv_data_simple(cv_text: str) -> CVExtractedData:
#     """
#     Simple regex-based extraction.
#     No AI needed!
#     """
#     print("\n🔍 Parsing CV with simple parser...")
    
#     name = extract_name(cv_text)
#     email = extract_email(cv_text)
#     phone = extract_phone(cv_text)
#     skills = extract_skills(cv_text)
#     education = extract_education(cv_text)
#     experience = extract_experience(cv_text)
    
#     # Guess total experience
#     total_exp = None
#     if experience:
#         # Extract years from all durations
#         years = 0
#         for exp in experience:
#             if exp.duration:
#                 year_match = re.search(r'(\d+)\s*year', exp.duration, re.IGNORECASE)
#                 if year_match:
#                     years += int(year_match.group(1))
#         total_exp = float(years) if years > 0 else None
    
#     extracted = CVExtractedData(
#         name=name,
#         email=email,
#         phone=phone,
#         total_exp=total_exp,
#         skills=skills,
#         education=education,
#         experience=experience
#     )
    
#     print(f"✅ Extracted:")
#     print(f"   Name: {name}")
#     print(f"   Email: {email}")
#     print(f"   Phone: {phone}")
#     print(f"   Skills: {len(skills)} found")
#     print(f"   Education: {len(education) if education else 0} found")
#     print(f"   Experience: {len(experience) if experience else 0} found")
    
#     return extracted


# def process_cv_file(filename: str, file_bytes: bytes) -> tuple[str, CVExtractedData]:
#     """
#     Main function - works without AI!
#     """
#     # Extract text
#     raw_text = extract_text_from_file(filename, file_bytes)
    
#     # Extract data with simple parser
#     extracted_data = extract_cv_data_simple(raw_text)
    
#     return raw_text, extracted_data



"""
cv_extractor.py - FIXED VERSION (NULL characters removed)
─────────────────────────────────────────────
✅ PostgreSQL compatible - removes NULL bytes
"""

import re
import pdfplumber
import docx
from schemas import CVExtractedData, EducationItem, ExperienceItem


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """PDF থেকে text বের করে"""
    import io
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """DOCX থেকে text বের করে"""
    import io
    doc = docx.Document(io.BytesIO(file_bytes))
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return text.strip()


def extract_text_from_file(filename: str, file_bytes: bytes) -> str:
    """File type check করে text extract করে"""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError("শুধু PDF বা DOCX support করা হয়")


def clean_text(text: str) -> str:
    """
    Remove NULL characters যা PostgreSQL accept করে না।
    
    NULL characters (0x00) PDF parsing এ মাঝে মাঝে আসে।
    PostgreSQL এগুলো string এ allow করে না।
    """
    if not text:
        return text
    
    # Remove NULL bytes (এটাই main সমস্যা)
    text = text.replace('\x00', '')
    text = text.replace('\u0000', '')
    
    # Optional: অন্য problematic control characters remove
    # যদি আরো সমস্যা হয় তাহলে এই লাইনটা uncomment করো
    # text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    
    return text


def extract_email(text: str) -> str:
    """Email বের করে"""
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    match = re.search(pattern, text)
    return match.group(0) if match else None


def extract_phone(text: str) -> str:
    """Phone number বের করে"""
    patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'\d{10,}',
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            return match.group(0)
    return None


def extract_name(text: str) -> str:
    """Name extraction - CV এর প্রথম কয়েক লাইন থেকে"""
    lines = text.split('\n')[:5]
    for line in lines:
        line = line.strip()
        if len(line) > 2 and len(line) < 50:
            if '@' not in line and not re.search(r'\d{3,}', line):
                if re.match(r'^[A-Za-z\s.]+$', line):
                    return line
    return None


def extract_skills(text: str) -> list:
    """Common skills খুঁজে বের করে"""
    skills_keywords = [
        # Programming
        'Python', 'Java', 'JavaScript', 'C++', 'C#', 'PHP', 'Ruby', 'Go', 'Rust',
        'TypeScript', 'Swift', 'Kotlin', 'R', 'MATLAB', 'Scala',
        
        # Web
        'HTML', 'CSS', 'React', 'Angular', 'Vue', 'Node.js', 'Express',
        'Django', 'Flask', 'FastAPI', 'Spring', 'ASP.NET', 'Laravel',
        
        # Database
        'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle',
        'SQLite', 'Cassandra', 'DynamoDB',
        
        # Tools
        'Git', 'Docker', 'Kubernetes', 'Jenkins', 'AWS', 'Azure', 'GCP',
        'Linux', 'REST API', 'GraphQL', 'Postman',
        
        # Other
        'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch',
        'Data Analysis', 'Excel', 'Power BI', 'Tableau',
    ]
    
    found_skills = []
    text_upper = text.upper()
    
    for skill in skills_keywords:
        if skill.upper() in text_upper:
            if skill not in found_skills:
                found_skills.append(skill)
    
    return found_skills


def extract_education(text: str) -> list:
    """Education খুঁজে বের করে"""
    education = []
    
    degree_patterns = [
        r'(B\.?Sc\.?|Bachelor|BSc|BS)\s+(in\s+)?([A-Za-z\s]+)',
        r'(M\.?Sc\.?|Master|MSc|MS)\s+(in\s+)?([A-Za-z\s]+)',
        r'(PhD|Ph\.D\.?|Doctorate)\s+(in\s+)?([A-Za-z\s]+)',
    ]
    
    for pattern in degree_patterns:
        matches = re.finditer(pattern, text, re.IGNORECASE)
        for match in matches:
            degree = match.group(0).strip()
            
            year_match = re.search(r'(19|20)\d{2}', text[max(0, match.start()-50):match.end()+50])
            year = year_match.group(0) if year_match else None
            
            education.append(EducationItem(
                degree=degree,
                institution=None,
                year=year
            ))
    
    return education if education else None


def extract_experience(text: str) -> list:
    """Work experience খুঁজে বের করে"""
    experience = []
    
    job_titles = [
        'Developer', 'Engineer', 'Manager', 'Designer', 'Analyst',
        'Consultant', 'Specialist', 'Architect', 'Lead', 'Senior',
        'Junior', 'Intern', 'Associate', 'Director', 'CEO', 'CTO'
    ]
    
    for title in job_titles:
        pattern = rf'\b({title})\b'
        matches = re.finditer(pattern, text, re.IGNORECASE)
        
        for match in matches:
            duration_match = re.search(r'(\d+)\s*(year|yr|month|mo)', 
                                      text[max(0, match.start()-100):match.end()+100], 
                                      re.IGNORECASE)
            duration = duration_match.group(0) if duration_match else None
            
            experience.append(ExperienceItem(
                role=match.group(0),
                company=None,
                duration=duration
            ))
    
    seen = set()
    unique_exp = []
    for exp in experience:
        if exp.role not in seen:
            seen.add(exp.role)
            unique_exp.append(exp)
    
    return unique_exp if unique_exp else None


def extract_cv_data_simple(cv_text: str) -> CVExtractedData:
    """Simple regex-based extraction"""
    print("\n🔍 Parsing CV with simple parser...")
    
    name = extract_name(cv_text)
    email = extract_email(cv_text)
    phone = extract_phone(cv_text)
    skills = extract_skills(cv_text)
    education = extract_education(cv_text)
    experience = extract_experience(cv_text)
    
    total_exp = None
    if experience:
        years = 0
        for exp in experience:
            if exp.duration:
                year_match = re.search(r'(\d+)\s*year', exp.duration, re.IGNORECASE)
                if year_match:
                    years += int(year_match.group(1))
        total_exp = float(years) if years > 0 else None
    
    extracted = CVExtractedData(
        name=name,
        email=email,
        phone=phone,
        total_exp=total_exp,
        skills=skills,
        education=education,
        experience=experience
    )
    
    print(f"✅ Extracted:")
    print(f"   Name: {name}")
    print(f"   Email: {email}")
    print(f"   Phone: {phone}")
    print(f"   Skills: {len(skills)} found")
    print(f"   Education: {len(education) if education else 0} found")
    print(f"   Experience: {len(experience) if experience else 0} found")
    
    return extracted


def process_cv_file(filename: str, file_bytes: bytes) -> tuple[str, CVExtractedData]:
    """
    Main function - processes CV file and returns clean data
    
    ✅ FIXED: Removes NULL characters before saving to PostgreSQL
    """
    print("📁 File received:", filename)
    print("🔍 CV processing started...")
    
    # Step 1: Extract text from file
    raw_text = extract_text_from_file(filename, file_bytes)
    
    # Step 2: Clean text - Remove NULL characters
    # এটা খুব গুরুত্বপূর্ণ! PostgreSQL NULL characters accept করে না।
    raw_text = clean_text(raw_text)
    
    # Step 3: Parse with simple parser
    extracted_data = extract_cv_data_simple(raw_text)

    return raw_text, extracted_data